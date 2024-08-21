# -*- coding: utf-8 -*-
"""
Created on Wed Aug 21 10:35:34 2024

@author: User
"""
#%%

import time
from tkinter import Tk, filedialog, Button, Label, Entry, StringVar, Scrollbar, Frame, messagebox
from docx import Document



class FileChangeHandler:
    """This class handles the update of a base document with a lsit of 
    replacement docs"""
    
    def __init__(self, input_doc_path, replacement_docs):
        self.input_doc_path = input_doc_path
        self.replacement_docs = replacement_docs

    def update_output_document(self):
        try:
            print("Updating base document.")
            doc = self.try_open_document(self.input_doc_path)
            doc = replace_multiple_sections(doc, self.replacement_docs)
            doc.save(self.input_doc_path)  # Save the changes to the same document
            print("Base document updated.")
            result_var.set("Base document updated successfully.")
            
        #Catches errors in case file is locked    
        except IOError as e:
            if "locked" in str(e).lower():
                print("Error: File is locked. Please close the file and try again.")
                result_var.set("Error: File is locked. Please close the file and try again.")
                messagebox.showerror("File Locked", "The file is currently locked. Please close it and try again.")
            else:
                print(f"Error updating document: {e}")
                result_var.set(f"Error updating document: {e}")
                messagebox.showerror("Update Error", f"Error updating document: {e}")

    def try_open_document(self, path, retries=3, delay=1):
        """Try to open a document with retries if it is locked."""
        for attempt in range(retries):
            try:
                return Document(path)
            except IOError as e:
                if "locked" in str(e).lower():
                    print(f"File is locked. Retry {attempt + 1}/{retries}...")
                    time.sleep(delay)
                else:
                    raise e
        raise RuntimeError("Failed to open file after several attempts")

def replace_multiple_sections(doc, replacement_docs):

    for replacement_doc_path in replacement_docs:
        replacement_doc = Document(replacement_doc_path)
        replacement_headings = find_headings(replacement_doc)
        
        print(replacement_headings)
        for section_heading, section_content in replacement_headings.items():
            section_start = None
            section_end = None
            
            # Find the start of the section in the base document
            for i, para in enumerate(doc.paragraphs):
                if para.style.name.startswith('Heading') and para.text == section_heading:
                    section_start = i
                    print("WORKING")
                    break
            
            if section_start is None:
                print(f"Section '{section_heading}' not found in base document")
                continue

            # Find the end of the section in the base document
            for j in range(section_start + 1, len(doc.paragraphs)):
                if doc.paragraphs[j].style.name.startswith('Heading'):
                    section_end = j
                    break
            else:
                section_end = len(doc.paragraphs)

            # Remove existing paragraphs in the section
            for _ in range(section_start + 1, section_end):
                p = doc.paragraphs[section_start + 1]
                p._element.getparent().remove(p._element)

            # Insert new paragraphs from the replacement document section
            for element in reversed(section_content):
                doc.element.body.insert(section_start + 1, element)

    return doc

def find_headings(doc):
    headings = {}
    
    current_heading = None
    current_content = []
    
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            if current_heading is not None:
                headings[current_heading] = current_content
            current_heading = para.text
            current_content = [para._element]
        else:
            if current_heading is not None:
                current_content.append(para._element)

    if current_heading is not None:
        headings[current_heading] = current_content
    
    return headings

def browse_file(entry):
    filename = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
    entry.delete(0, 'end')
    entry.insert(0, filename)

def add_replacement_entry():
    global replacement_entries
    row = len(replacement_entries) + 1
    file_label = Label(replacement_frame, text=f"Replacement Document {row}")
    file_label.grid(row=row, column=0, padx=10, pady=5)
    file_entry = Entry(replacement_frame, width=50)
    file_entry.grid(row=row, column=1, padx=10, pady=5)
    browse_button = Button(replacement_frame, text="Browse", command=lambda: browse_file(file_entry))
    browse_button.grid(row=row, column=2, padx=10, pady=5)

    replacement_entries.append(file_entry)

def update_document():
    input_doc_path = input_file_entry.get()
    
    # Read replacement paths
    replacement_docs = [entry.get() for entry in replacement_entries if entry.get()]
    print(replacement_docs)
    
    global file_change_handler
    file_change_handler = FileChangeHandler(input_doc_path, replacement_docs)
    file_change_handler.update_output_document()

# GUI setup code would go here...

# Setup GUI
root = Tk()
root.title("MUFG Collab")

# Set colors
bg_color = 'white'
fg_color = 'red'
btn_bg_color = 'red'
btn_fg_color = 'white'

# Apply colors to root window
root.configure(bg=bg_color)

# Title
title_label = Label(root, text="MUFG Collab", font=("Helvetica", 18, "bold"), bg=bg_color, fg=fg_color)
title_label.grid(row=0, column=0, columnspan=3, padx=10, pady=10)

# Input Document
Label(root, text="Base Document", bg=bg_color, fg=fg_color, font=("Helvetica", 12)).grid(row=1, column=0, padx=10, pady=10, sticky='w')
input_file_entry = Entry(root, width=50, font=("Helvetica", 12))
input_file_entry.grid(row=1, column=1, padx=10, pady=10)
Button(root, text="Browse", command=lambda: browse_file(input_file_entry), bg=btn_bg_color, fg=btn_fg_color, font=("Helvetica", 12, "bold")).grid(row=1, column=2, padx=10, pady=10)

# Frame for replacement documents
replacement_frame = Frame(root, bg=bg_color)
replacement_frame.grid(row=2, column=0, columnspan=3, padx=10, pady=10, sticky='nsew')
replacement_entries = []

# Scrollbar for the replacement documents frame
scrollbar = Scrollbar(replacement_frame, orient="vertical")
scrollbar.grid(row=0, column=3, rowspan=10, sticky='ns')

# Add initial replacement entry
add_replacement_entry_button = Button(root, text="Add Replacement Document", command=add_replacement_entry, bg=btn_bg_color, fg=btn_fg_color, font=("Helvetica", 12, "bold"))
add_replacement_entry_button.grid(row=3, column=0, columnspan=3, padx=10, pady=10)

# Update Document button
Button(root, text="Update Document", command=update_document, bg=btn_bg_color, fg=btn_fg_color, font=("Helvetica", 12, "bold")).grid(row=4, column=0, columnspan=3, padx=10, pady=10)

result_var = StringVar()
result_label = Label(root, textvariable=result_var, bg=bg_color, fg=fg_color, font=("Helvetica", 12))
result_label.grid(row=5, column=0, columnspan=3, padx=10, pady=10)

# Start Tkinter event loop
root.mainloop()
